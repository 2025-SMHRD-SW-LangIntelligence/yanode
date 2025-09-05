# rag_agent/ops_dooray.py

import os, re, unicodedata, time
from concurrent.futures import ThreadPoolExecutor, as_completed

# 선택: PyMuPDF(있으면 샘플 텍스트가 더 빠름)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore

import pdfplumber
import pytesseract
from PIL import Image, ImageOps

from rag_agent.indexer import (
    hybrid_search, GLOBAL_INDEX, ensure_index, CHECKLIST_FOLDER_IDS, _ext_lower
)
from rag_agent.previewers import (
    _preview_pdf_by_item, _preview_hwp_by_item, _preview_docx_by_item,
    _preview_xlsx_by_item, _preview_txt_by_item, _preview_ppt_like_by_item
)
from rag_agent.downloader import _download_to_tmp, _extract_pdf_text, _ocr_pdf
from rag_agent.api import list_personal_drives, search_in_root

# ===== 설정 =====
_UI_HOST = "smhrd.dooray.com"     # 하드코딩
FALLBACK_TIME_BUDGET_SEC = 12.0         # 본문 폴백 총 예산
CHEAP_POOL_WORKERS = 4                  # 경량 스니핑 병렬도
DEEP_PDF_MAX = 8                        # 깊은 추출 PDF 개수(힌트 없을 때)
DEEP_PDF_MAX_HINT = 4                   # 깊은 추출 PDF 개수(힌트 있을 때)
PDF_TEXT_SAMPLES = 10                   # PDF 텍스트 샘플 페이지 수(앞/중/끝 고르게)
PDF_OCR_SAMPLES  = 5                    # PDF OCR 샘플 페이지 수(스캔일 때만)
OCR_PAGE_STRIDE  = 2                    # 깊은 OCR stride

def _file_url(drive_id: str, file_id: str) -> str:
    return f"https://{_UI_HOST}/drive/{drive_id}/{file_id}"

# ===== 스코프 =====
def _is_in_scope(it: dict) -> bool:
    allowed = set(GLOBAL_INDEX.get("roots") or [])
    if not allowed:
        return True
    rid = it.get("root_id") or it.get("rootId") or it.get("topParentId")
    if rid and rid in allowed:
        return True
    ancestors = it.get("_ancestors") or []
    if any(a in allowed for a in ancestors):
        return True
    root_paths = GLOBAL_INDEX.get("root_paths") or []
    pth = it.get("_path") or ""
    if any(pth.startswith(rp) for rp in root_paths):
        return True
    return False

def _ensure_index_keep_scope():
    """현재 설정된 roots(체크리스트)를 보존하면서 인덱스를 보장한다.
       전체 선택(=제한 없음)일 때는 roots를 None으로 명확히 비운다.
    """
    # 전체 선택으로 취급할 패턴들
    ALL_SENTINELS = (None, [], ["root"], ["*"], ["ALL"])

    if CHECKLIST_FOLDER_IDS in ALL_SENTINELS:
        # ➜ 제한 해제: 전체 드라이브 인덱스
        ensure_index(None, force=False)
        GLOBAL_INDEX["roots"] = None
        GLOBAL_INDEX["root_paths"] = []
    else:
        # ➜ 특정 폴더만 인덱스
        ensure_index(CHECKLIST_FOLDER_IDS, force=False)


# ===== 프리뷰 디스패처 =====
def resolve_file_handler(item: dict) -> str:
    ext = (item.get("ext") or "").lower()
    if ext == ".pdf":  return _preview_pdf_by_item(item)
    if ext in (".hwp", ".hwpx"): return _preview_hwp_by_item(item)
    if ext == ".docx": return _preview_docx_by_item(item)
    if ext == ".xlsx": return _preview_xlsx_by_item(item)
    if ext == ".txt":  return _preview_txt_by_item(item)
    if ext in (".ppt", ".pptx"): return _preview_ppt_like_by_item(item)
    return f"❓ 지원하지 않는 형식: {ext} ({item.get('name','')})"

# ===== 질의 정규화 =====
_EXT_MAP = {
    "pdf": (".pdf",), "hwp": (".hwp", ".hwpx"),
    "ppt": (".ppt", ".pptx"), "pptx": (".pptx",),
    "docx": (".docx",), "word": (".docx",),
    "txt": (".txt",), "excel": (".xlsx", ".xls"), "xlsx": (".xlsx",)
}
def _extract_ext_filter(q: str):
    ql = (q or "").lower()
    want = []
    for k, v in _EXT_MAP.items():
        if re.search(rf"\b{k}\b", ql): want.extend(v)
    return tuple(sorted(set(want))) if want else None

def _strip_preview_header(s: str) -> str:
    return re.sub(r"^📄 \*.*?\* 미리보기.*?:\s*", "", (s or "").strip(), flags=re.S)

def _norm_kor(s: str) -> str:
    s = unicodedata.normalize("NFKC", s or "").lower()
    s = re.sub(r"\s+", "", s)
    return "".join(ch for ch in s if ch.isalnum() or '\uAC00' <= ch <= '\uD7A3')

def _ngrams(t: str, n: int) -> list[str]:
    if n <= 1: return [t]
    return [t[i:i+n] for i in range(len(t)-n+1)] if len(t) >= n else []

def _make_query_terms(q: str) -> list[str]:
    if not q: return []
    base = q.strip()
    full_norm = _norm_kor(base)
    toks = [_norm_kor(t) for t in re.split(r"\s+", base) if t]
    out = set([t for t in toks if len(t) >= 2])
    if len(full_norm) >= 2: out.add(full_norm)
    out.update(_ngrams(full_norm, 2)); out.update(_ngrams(full_norm, 3))
    return list(out)

def _score_in_text(text: str, toks: list[str]) -> int:
    tl_raw = (text or "").lower()
    tl_norm = _norm_kor(tl_raw)
    s = 0
    for t in toks:
        tr, tn = (t or "").lower(), _norm_kor(t or "")
        if tr: s += tl_raw.count(tr)
        if tn: s += 2 * tl_norm.count(tn)
    return s

def _text_contains_query(item: dict, terms: list[str]) -> tuple[int, str]:
    try:
        prev = resolve_file_handler(item)
        body = _strip_preview_header(prev)
        return _score_in_text(body, terms), body
    except Exception:
        return 0, ""
    

# ===== PDF 경량 샘플링 (텍스트 → 부족시 소량 OCR) =====
def _quantized_indices(n: int, k: int) -> list[int]:
    """0..n-1 구간을 k개로 고르게 샘플."""
    if n <= 0 or k <= 0: return []
    if k >= n: return list(range(n))
    return sorted(set(int(round(i * (n - 1) / (k - 1))) for i in range(k)))


def _take_uniform_by_ext(name_scores, exts_tuple, cap):
    bucket = [it for _, it in name_scores if (it.get("ext") or "").lower() in exts_tuple]
    if not bucket or cap <= 0:
        return []
    if len(bucket) <= cap:
        return bucket
    idxs = _quantized_indices(len(bucket), cap)
    return [bucket[i] for i in idxs]



def _cheap_pdf_sample_text(src: str, k: int = PDF_TEXT_SAMPLES) -> str:
    """
    PDF 텍스트 경량 샘플링:
    - PyMuPDF(fitz)가 있으면 빠르게 페이지 텍스트 추출(앞/중/끝 균등 샘플)
    - 없으면 pdfplumber로 '전체 페이지 수'를 보고 균등 샘플해 extract_text()
    """
    # 1) PyMuPDF 경로
    if fitz:
        try:
            doc = fitz.open(src)
            n = len(doc)
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for i in idxs:
                try:
                    t = (doc[i].get_text("text") or "").strip()
                    if t:
                        parts.append(f"--- [Page {i+1}] ---\n{t}")
                except Exception:
                    continue
            doc.close()
            return "\n\n".join(parts).strip()
        except Exception:
            pass  # fitz 실패 시 pdfplumber 폴백으로 이동

    # 2) pdfplumber 폴백 (앞/중/끝 균등 샘플)
    try:
        with pdfplumber.open(src) as pdf:
            n = len(pdf.pages)
            if n <= 0:
                return ""
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for i in idxs:
                try:
                    t = (pdf.pages[i].extract_text() or "").strip()
                    if t:
                        parts.append(f"--- [Page {i+1}] ---\n{t}")
                except Exception:
                    continue
            return "\n\n".join(parts).strip()
    except Exception:
        return ""

    return ""


def _cheap_pdf_sample_ocr(src: str, k: int = PDF_OCR_SAMPLES, dpi: int = 220,
                          lang_primary: str = "kor+eng", lang_fallback: str = "eng") -> str:
    """스캔 PDF 대비: 소량 페이지만 OCR(앞/중/끝), 빠름."""
    try:
        with pdfplumber.open(src) as pdf:
            n = len(pdf.pages)
            if n == 0: return ""
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for idx in idxs:
                try:
                    pil = pdf.pages[idx].to_image(resolution=dpi).original
                    g = ImageOps.autocontrast(pil.convert("L"))
                    cfg = "--oem 3 --psm 6"
                    txt = pytesseract.image_to_string(g, lang=lang_primary, config=cfg) or ""
                    if not txt:
                        txt = pytesseract.image_to_string(g, lang=lang_fallback, config=cfg) or ""
                    txt = txt.strip()
                    if txt:
                        parts.append(f"--- [OCR Page {idx+1}] ---\n{txt}")
                except Exception:
                    continue
            return "\n\n".join(parts).strip()
    except Exception:
        return ""
    
def _cheap_pdf_body(it: dict) -> str:
    """PDF 경량 본문: 문서 전체 범위를 앞/중/끝 균등 샘플링 → 부족하면 소량 OCR."""
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives:
            return ""
        drive_id = drives[0]["id"]

    try:
        src = _download_to_tmp(it["id"], ".pdf", drive_id)

        # 1) 텍스트 샘플(문서 전체에서 균등 분할)
        txt = _cheap_pdf_sample_text(src, k=PDF_TEXT_SAMPLES) or ""
        if len(txt.strip()) >= 30:
            return txt

        # 2) 스캔/저품질 대비: 소량 페이지만 OCR(앞/중/끝)
        ocr = _cheap_pdf_sample_ocr(src, k=PDF_OCR_SAMPLES) or ""
        if ocr:
            return (txt + "\n\n" + ocr).strip() if txt else ocr

        return txt
    except Exception:
        return ""
    


def _cheap_body_sniff(item: dict) -> str:
    """경량 1차 점수화(절대 전체 OCR 금지)."""
    ext = (item.get("ext") or "").lower()

    # ✅ PDF는 전용 경량 바디 함수로 위임(앞/중/끝 샘플 + 필요시 소량 OCR)
    if ext == ".pdf":
        return _cheap_pdf_body(item)

    # ⬇️ PDF 이외만 기존 로직 유지
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives:
            return ""
        drive_id = drives[0]["id"]

    try:
        if ext == ".txt":
            src = _download_to_tmp(item["id"], ".txt", drive_id)
            return open(src, "r", encoding="utf-8", errors="ignore").read()

        if ext == ".docx":
            src = _download_to_tmp(item["id"], ".docx", drive_id)
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                docs = Docx2txtLoader(src).load()
                return "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
            except Exception:
                try:
                    from docx import Document as _Docx
                    d = _Docx(src)
                    return "\n".join(p.text.strip() for p in d.paragraphs if p.text and p.text.strip())
                except Exception:
                    return ""

        # HWP/PPT/XLSX 등은 프리뷰 본문 사용
        _, body = _text_contains_query(item, [])
        return body
    except Exception:
        return ""

# 간단 캐시(파일ID+확장자 키)
def _sniff_with_cache(it: dict) -> str:
    key = f"{it.get('id','')}::{(it.get('ext') or '').lower()}"
    if not hasattr(_sniff_with_cache, "_cache"):
        _sniff_with_cache._cache = {}
    cache = _sniff_with_cache._cache  # type: ignore
    if key in cache: return cache[key]
    body = _cheap_body_sniff(it)
    cache[key] = body
    if len(cache) > 256:
        cache.pop(next(iter(cache)))
    return body

# ===== 깊은 폴백(긴 본문 / OCR 허용) =====
def _fulltext_for_scoring(item: dict, pdf_pages: int = 15) -> str:
    ext = (item.get("ext") or "").lower()
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives: return ""
        drive_id = drives[0]["id"]
    try:
        if ext == ".pdf":
            src = _download_to_tmp(item["id"], ".pdf", drive_id)
            txt = _extract_pdf_text(src, max_pages=pdf_pages)
            if not txt or len(txt.strip()) < 50:
                txt = _ocr_pdf(src, max_pages=pdf_pages, page_stride=OCR_PAGE_STRIDE)
            return txt or ""
        if ext in (".hwp", ".hwpx"):
            try:
                from langchain_teddynote.document_loaders import HWPLoader
            except Exception:
                return ""
            src = _download_to_tmp(item["id"], ext, drive_id)
            try:
                docs = HWPLoader(src).load()
                return "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
            except Exception:
                return ""
        if ext == ".docx":
            src = _download_to_tmp(item["id"], ".docx", drive_id)
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                docs = Docx2txtLoader(src).load()
                txt = "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
                if txt: return txt
            except Exception:
                pass
            try:
                from docx import Document as _Docx
                d = _Docx(src)
                return "\n".join(p.text.strip() for p in d.paragraphs if p.text and p.text.strip())
            except Exception:
                return ""
        if ext == ".txt":
            src = _download_to_tmp(item["id"], ".txt", drive_id)
            return open(src, "r", encoding="utf-8", errors="ignore").read()
    except Exception:
        return ""
    return ""

# ===== 전역(루트) 후보 폴백 =====
def _global_root_candidates(query: str, exts: tuple[str, ...] | None, cap: int = 30) -> list[dict]:
    drives = list_personal_drives()
    if not drives: return []
    drive_id = drives[0]["id"]
    hits = search_in_root(drive_id, query) or []
    out = []
    for it in hits:
        if it.get("type") != "file":  # 폴더는 제외(속도)
            continue
        ext = _ext_lower(it.get("name", ""))
        if exts and ext not in exts: continue
        out.append({"id": it["id"], "name": it["name"], "type": "file", "_path": "/", "ext": ext})
        if len(out) >= cap: break
    return out

# =========================================
# 메인
# =========================================
# rag_agent/ops_dooray.py

import os, re, unicodedata, time
from concurrent.futures import ThreadPoolExecutor, as_completed

# 선택: PyMuPDF(있으면 샘플 텍스트가 더 빠름)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None  # type: ignore

import pdfplumber
import pytesseract
from PIL import Image, ImageOps

from rag_agent.indexer import (
    hybrid_search, GLOBAL_INDEX, ensure_index, CHECKLIST_FOLDER_IDS, _ext_lower
)
from rag_agent.previewers import (
    _preview_pdf_by_item, _preview_hwp_by_item, _preview_docx_by_item,
    _preview_xlsx_by_item, _preview_txt_by_item, _preview_ppt_like_by_item
)
from rag_agent.downloader import _download_to_tmp, _extract_pdf_text, _ocr_pdf
from rag_agent.api import list_personal_drives, search_in_root

# ===== 설정 =====
_UI_HOST = "smhrd.dooray.com"     # 하드코딩
FALLBACK_TIME_BUDGET_SEC = 12.0         # 본문 폴백 총 예산
CHEAP_POOL_WORKERS = 4                  # 경량 스니핑 병렬도
DEEP_PDF_MAX = 8                        # 깊은 추출 PDF 개수(힌트 없을 때)
DEEP_PDF_MAX_HINT = 4                   # 깊은 추출 PDF 개수(힌트 있을 때)
PDF_TEXT_SAMPLES = 10                   # PDF 텍스트 샘플 페이지 수(앞/중/끝 고르게)
PDF_OCR_SAMPLES  = 5                    # PDF OCR 샘플 페이지 수(스캔일 때만)
OCR_PAGE_STRIDE  = 2                    # 깊은 OCR stride

def _file_url(drive_id: str, file_id: str) -> str:
    return f"https://{_UI_HOST}/drive/{drive_id}/{file_id}"

# ===== 스코프 =====
def _is_in_scope(it: dict) -> bool:
    allowed = set(GLOBAL_INDEX.get("roots") or [])
    if not allowed:
        return True
    rid = it.get("root_id") or it.get("rootId") or it.get("topParentId")
    if rid and rid in allowed:
        return True
    ancestors = it.get("_ancestors") or []
    if any(a in allowed for a in ancestors):
        return True
    root_paths = GLOBAL_INDEX.get("root_paths") or []
    pth = it.get("_path") or ""
    if any(pth.startswith(rp) for rp in root_paths):
        return True
    return False

def _ensure_index_keep_scope():
    """현재 설정된 roots(체크리스트)를 보존하면서 인덱스를 보장한다."""
    current_roots = GLOBAL_INDEX.get("roots")
    if CHECKLIST_FOLDER_IDS:  # 외부에서 강제 지정된 체크리스트가 있으면 그걸 사용
        ensure_index(CHECKLIST_FOLDER_IDS, force=False)
    else:
        # 현재 설정된 roots를 그대로 유지(없으면 None)
        ensure_index(list(current_roots) if current_roots else None, force=False)


# ===== 프리뷰 디스패처 =====
def resolve_file_handler(item: dict) -> str:
    ext = (item.get("ext") or "").lower()
    if ext == ".pdf":  return _preview_pdf_by_item(item)
    if ext in (".hwp", ".hwpx"): return _preview_hwp_by_item(item)
    if ext == ".docx": return _preview_docx_by_item(item)
    if ext == ".xlsx": return _preview_xlsx_by_item(item)
    if ext == ".txt":  return _preview_txt_by_item(item)
    if ext in (".ppt", ".pptx"): return _preview_ppt_like_by_item(item)
    return f"❓ 지원하지 않는 형식: {ext} ({item.get('name','')})"

# ===== 질의 정규화 =====
_EXT_MAP = {
    "pdf": (".pdf",), "hwp": (".hwp", ".hwpx"),
    "ppt": (".ppt", ".pptx"), "pptx": (".pptx",),
    "docx": (".docx",), "word": (".docx",),
    "txt": (".txt",), "excel": (".xlsx", ".xls"), "xlsx": (".xlsx",)
}
def _extract_ext_filter(q: str):
    ql = (q or "").lower()
    want = []
    for k, v in _EXT_MAP.items():
        if re.search(rf"\b{k}\b", ql): want.extend(v)
    return tuple(sorted(set(want))) if want else None

def _strip_preview_header(s: str) -> str:
    return re.sub(r"^📄 \*.*?\* 미리보기.*?:\s*", "", (s or "").strip(), flags=re.S)

def _norm_kor(s: str) -> str:
    s = unicodedata.normalize("NFKC", s or "").lower()
    s = re.sub(r"\s+", "", s)
    return "".join(ch for ch in s if ch.isalnum() or '\uAC00' <= ch <= '\uD7A3')

def _ngrams(t: str, n: int) -> list[str]:
    if n <= 1: return [t]
    return [t[i:i+n] for i in range(len(t)-n+1)] if len(t) >= n else []

def _make_query_terms(q: str) -> list[str]:
    if not q: return []
    base = q.strip()
    full_norm = _norm_kor(base)
    toks = [_norm_kor(t) for t in re.split(r"\s+", base) if t]
    out = set([t for t in toks if len(t) >= 2])
    if len(full_norm) >= 2: out.add(full_norm)
    out.update(_ngrams(full_norm, 2)); out.update(_ngrams(full_norm, 3))
    return list(out)

def _score_in_text(text: str, toks: list[str]) -> int:
    tl_raw = (text or "").lower()
    tl_norm = _norm_kor(tl_raw)
    s = 0
    for t in toks:
        tr, tn = (t or "").lower(), _norm_kor(t or "")
        if tr: s += tl_raw.count(tr)
        if tn: s += 2 * tl_norm.count(tn)
    return s

def _text_contains_query(item: dict, terms: list[str]) -> tuple[int, str]:
    try:
        prev = resolve_file_handler(item)
        body = _strip_preview_header(prev)
        return _score_in_text(body, terms), body
    except Exception:
        return 0, ""
    

# ===== PDF 경량 샘플링 (텍스트 → 부족시 소량 OCR) =====
def _quantized_indices(n: int, k: int) -> list[int]:
    """0..n-1 구간을 k개로 고르게 샘플."""
    if n <= 0 or k <= 0: return []
    if k >= n: return list(range(n))
    return sorted(set(int(round(i * (n - 1) / (k - 1))) for i in range(k)))


def _take_uniform_by_ext(name_scores, exts_tuple, cap):
    bucket = [it for _, it in name_scores if (it.get("ext") or "").lower() in exts_tuple]
    if not bucket or cap <= 0:
        return []
    if len(bucket) <= cap:
        return bucket
    idxs = _quantized_indices(len(bucket), cap)
    return [bucket[i] for i in idxs]



def _cheap_pdf_sample_text(src: str, k: int = PDF_TEXT_SAMPLES) -> str:
    """
    PDF 텍스트 경량 샘플링:
    - PyMuPDF(fitz)가 있으면 빠르게 페이지 텍스트 추출(앞/중/끝 균등 샘플)
    - 없으면 pdfplumber로 '전체 페이지 수'를 보고 균등 샘플해 extract_text()
    """
    # 1) PyMuPDF 경로
    if fitz:
        try:
            doc = fitz.open(src)
            n = len(doc)
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for i in idxs:
                try:
                    t = (doc[i].get_text("text") or "").strip()
                    if t:
                        parts.append(f"--- [Page {i+1}] ---\n{t}")
                except Exception:
                    continue
            doc.close()
            return "\n\n".join(parts).strip()
        except Exception:
            pass  # fitz 실패 시 pdfplumber 폴백으로 이동

    # 2) pdfplumber 폴백 (앞/중/끝 균등 샘플)
    try:
        with pdfplumber.open(src) as pdf:
            n = len(pdf.pages)
            if n <= 0:
                return ""
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for i in idxs:
                try:
                    t = (pdf.pages[i].extract_text() or "").strip()
                    if t:
                        parts.append(f"--- [Page {i+1}] ---\n{t}")
                except Exception:
                    continue
            return "\n\n".join(parts).strip()
    except Exception:
        return ""

    return ""


def _cheap_pdf_sample_ocr(src: str, k: int = PDF_OCR_SAMPLES, dpi: int = 220,
                          lang_primary: str = "kor+eng", lang_fallback: str = "eng") -> str:
    """스캔 PDF 대비: 소량 페이지만 OCR(앞/중/끝), 빠름."""
    try:
        with pdfplumber.open(src) as pdf:
            n = len(pdf.pages)
            if n == 0: return ""
            idxs = _quantized_indices(n, min(k, n))
            parts = []
            for idx in idxs:
                try:
                    pil = pdf.pages[idx].to_image(resolution=dpi).original
                    g = ImageOps.autocontrast(pil.convert("L"))
                    cfg = "--oem 3 --psm 6"
                    txt = pytesseract.image_to_string(g, lang=lang_primary, config=cfg) or ""
                    if not txt:
                        txt = pytesseract.image_to_string(g, lang=lang_fallback, config=cfg) or ""
                    txt = txt.strip()
                    if txt:
                        parts.append(f"--- [OCR Page {idx+1}] ---\n{txt}")
                except Exception:
                    continue
            return "\n\n".join(parts).strip()
    except Exception:
        return ""
    
def _cheap_pdf_body(it: dict) -> str:
    """PDF 경량 본문: 문서 전체 범위를 앞/중/끝 균등 샘플링 → 부족하면 소량 OCR."""
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives:
            return ""
        drive_id = drives[0]["id"]

    try:
        src = _download_to_tmp(it["id"], ".pdf", drive_id)

        # 1) 텍스트 샘플(문서 전체에서 균등 분할)
        txt = _cheap_pdf_sample_text(src, k=PDF_TEXT_SAMPLES) or ""
        if len(txt.strip()) >= 30:
            return txt

        # 2) 스캔/저품질 대비: 소량 페이지만 OCR(앞/중/끝)
        ocr = _cheap_pdf_sample_ocr(src, k=PDF_OCR_SAMPLES) or ""
        if ocr:
            return (txt + "\n\n" + ocr).strip() if txt else ocr

        return txt
    except Exception:
        return ""
    


def _cheap_body_sniff(item: dict) -> str:
    """경량 1차 점수화(절대 전체 OCR 금지)."""
    ext = (item.get("ext") or "").lower()

    # ✅ PDF는 전용 경량 바디 함수로 위임(앞/중/끝 샘플 + 필요시 소량 OCR)
    if ext == ".pdf":
        return _cheap_pdf_body(item)

    # ⬇️ PDF 이외만 기존 로직 유지
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives:
            return ""
        drive_id = drives[0]["id"]

    try:
        if ext == ".txt":
            src = _download_to_tmp(item["id"], ".txt", drive_id)
            return open(src, "r", encoding="utf-8", errors="ignore").read()

        if ext == ".docx":
            src = _download_to_tmp(item["id"], ".docx", drive_id)
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                docs = Docx2txtLoader(src).load()
                return "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
            except Exception:
                try:
                    from docx import Document as _Docx
                    d = _Docx(src)
                    return "\n".join(p.text.strip() for p in d.paragraphs if p.text and p.text.strip())
                except Exception:
                    return ""

        # HWP/PPT/XLSX 등은 프리뷰 본문 사용
        _, body = _text_contains_query(item, [])
        return body
    except Exception:
        return ""

# 간단 캐시(파일ID+확장자 키)
def _sniff_with_cache(it: dict) -> str:
    key = f"{it.get('id','')}::{(it.get('ext') or '').lower()}"
    if not hasattr(_sniff_with_cache, "_cache"):
        _sniff_with_cache._cache = {}
    cache = _sniff_with_cache._cache  # type: ignore
    if key in cache: return cache[key]
    body = _cheap_body_sniff(it)
    cache[key] = body
    if len(cache) > 256:
        cache.pop(next(iter(cache)))
    return body

# ===== 깊은 폴백(긴 본문 / OCR 허용) =====
def _fulltext_for_scoring(item: dict, pdf_pages: int = 15) -> str:
    ext = (item.get("ext") or "").lower()
    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if not drives: return ""
        drive_id = drives[0]["id"]
    try:
        if ext == ".pdf":
            src = _download_to_tmp(item["id"], ".pdf", drive_id)
            txt = _extract_pdf_text(src, max_pages=pdf_pages)
            if not txt or len(txt.strip()) < 50:
                txt = _ocr_pdf(src, max_pages=pdf_pages, page_stride=OCR_PAGE_STRIDE)
            return txt or ""
        if ext in (".hwp", ".hwpx"):
            try:
                from langchain_teddynote.document_loaders import HWPLoader
            except Exception:
                return ""
            src = _download_to_tmp(item["id"], ext, drive_id)
            try:
                docs = HWPLoader(src).load()
                return "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
            except Exception:
                return ""
        if ext == ".docx":
            src = _download_to_tmp(item["id"], ".docx", drive_id)
            try:
                from langchain_community.document_loaders import Docx2txtLoader
                docs = Docx2txtLoader(src).load()
                txt = "\n".join((getattr(d, "page_content", "") or "") for d in docs).strip()
                if txt: return txt
            except Exception:
                pass
            try:
                from docx import Document as _Docx
                d = _Docx(src)
                return "\n".join(p.text.strip() for p in d.paragraphs if p.text and p.text.strip())
            except Exception:
                return ""
        if ext == ".txt":
            src = _download_to_tmp(item["id"], ".txt", drive_id)
            return open(src, "r", encoding="utf-8", errors="ignore").read()
    except Exception:
        return ""
    return ""

# ===== 전역(루트) 후보 폴백 =====
def _global_root_candidates(query: str, exts: tuple[str, ...] | None, cap: int = 30) -> list[dict]:
    drives = list_personal_drives()
    if not drives: return []
    drive_id = drives[0]["id"]
    hits = search_in_root(drive_id, query) or []
    out = []
    for it in hits:
        if it.get("type") != "file":  # 폴더는 제외(속도)
            continue
        ext = _ext_lower(it.get("name", ""))
        if exts and ext not in exts: continue
        out.append({"id": it["id"], "name": it["name"], "type": "file", "_path": "/", "ext": ext})
        if len(out) >= cap: break
    return out

# =========================================
# 메인
# =========================================
def dooray_auto_preview(query: str) -> str:
    """
    본문/OCR로 후보를 고르되, 항상 키=값 블록을 반환한다.
    성공: __TOP1_* 키들에 값 채움
    실패: __TOP1_* 는 '-' 로, __MSG__ 에 안내문 작성
    """
    exts = _extract_ext_filter(query)

    # 스코프(체크리스트) 보존
    _ensure_index_keep_scope()
    strict_scope = bool(GLOBAL_INDEX.get("roots"))

    # ---- 헬퍼: 표준 KV 블록 생성 ----
    def _kv(name="", path="", ext="", url="", preview="", msg=""):
        def _clip(s, n=1200):
            s = (s or "").strip()
            return (s[:n] + "…") if len(s) > n else s
        lines = [
            f"__TOP1_NAME__={name or '-'}",
            f"__TOP1_PATH__={path or '-'}",
            f"__TOP1_EXT__={(ext or '-')}",
            f"__TOP1_URL__={url or '-'}",
            f"__TOP1_PREVIEW__={_clip(preview) or '-'}",
        ]
        if msg:
            lines.append(f"__MSG__={msg}")
        return "\n".join(lines)

    # ---- 1) 파일명/경로 기반 ----
    items = hybrid_search(
        query,
        exts=exts or (".pdf", ".hwp", ".hwpx", ".ppt", ".pptx", ".docx", ".xlsx", ".txt"),
        limit=10
    )
    name_terms = _make_query_terms(query)

    def _name_hit(it: dict) -> bool:
        raw  = ((it.get("name","") + " " + it.get("_path","")) or "").lower()
        norm = _norm_kor(it.get("name","") + it.get("_path",""))
        return any(t in raw or t in norm for t in name_terms)

    items = [it for it in items if _is_in_scope(it) and _name_hit(it)]

    if items:
        best = items[0]
        try:
            preview = resolve_file_handler(best)
        except Exception:
            preview = "⚠️ 후보 파일 미리보기 중 오류가 발생했습니다."
        url = f"https://{_UI_HOST}/preview-pages/drives/{best['id']}"
        return _kv(best.get("name",""), best.get("_path","/"), best.get("ext",""),
                   url, preview, "")

    # ---- 2) 본문/의미 폴백(빠른 부정 포함) ----
    if not name_terms:
        return _kv(msg=f"'{query}' 관련 파일을 찾지 못했습니다.")

    t0 = time.perf_counter()
    ALL = [it for it in (GLOBAL_INDEX.get("items") or [])
           if _is_in_scope(it) and (not exts or (it.get("ext") in exts))]

    def _name_score(it):
        hay_raw = ((it.get("_path","") + " " + it.get("name","")) or "").lower()
        hay_norm = _norm_kor(hay_raw)
        s = 0
        for t in name_terms:
            s += hay_raw.count(t) + 2 * hay_norm.count(t)
        return s

    name_scores = [(_name_score(it), it) for it in ALL]
    name_scores.sort(key=lambda x: x[0], reverse=True)
    no_name_signal = (name_scores[0][0] == 0) if name_scores else True

    # 빠른 부정 모드
    if strict_scope and no_name_signal:
        TXT_CAP, HWP_CAP, PDF_CAP, PPT_CAP, XLSX_CAP = 12, 12, 16, 6, 4
        TIME_BUDGET = min(3.5, FALLBACK_TIME_BUDGET_SEC)
        USE_OCR = False
    else:
        TXT_CAP, HWP_CAP, PDF_CAP, PPT_CAP, XLSX_CAP = 40, 40, (120 if no_name_signal else 24), 12, 6
        TIME_BUDGET = FALLBACK_TIME_BUDGET_SEC
        USE_OCR = True

    def _take_uniform_by_ext_local(name_scores, exts_tuple, cap):
        bucket = [it for _, it in name_scores if (it.get("ext") or "").lower() in exts_tuple]
        if not bucket or cap <= 0:
            return []
        if len(bucket) <= cap:
            return bucket
        idxs = _quantized_indices(len(bucket), cap)
        return [bucket[i] for i in idxs]

    scope_items = (
        _take_uniform_by_ext_local(name_scores, (".txt",), TXT_CAP) +
        _take_uniform_by_ext_local(name_scores, (".hwp", ".hwpx"), HWP_CAP) +
        _take_uniform_by_ext_local(name_scores, (".pdf",), PDF_CAP) +
        _take_uniform_by_ext_local(name_scores, (".ppt", ".pptx"), PPT_CAP) +
        _take_uniform_by_ext_local(name_scores, (".xlsx",), XLSX_CAP)
    )

    drive_id = GLOBAL_INDEX.get("drive_id")
    if not drive_id:
        drives = list_personal_drives()
        if drives:
            drive_id = drives[0]["id"]

    def _quick_sniff(it: dict) -> str:
        ext = (it.get("ext") or "").lower()
        try:
            if ext == ".pdf":
                if not drive_id:
                    return ""
                src = _download_to_tmp(it["id"], ".pdf", drive_id)
                txt = _cheap_pdf_sample_text(src, k=min(6, PDF_TEXT_SAMPLES)) or ""
                if (not txt) and USE_OCR:
                    txt = _cheap_pdf_sample_ocr(src, k=3) or ""
                return txt
            if ext == ".txt":
                if not drive_id:
                    return ""
                src = _download_to_tmp(it["id"], ".txt", drive_id)
                return open(src, "r", encoding="utf-8", errors="ignore").read()
            if ext == ".docx":
                return _cheap_body_sniff(it)
            _, body = _text_contains_query(it, [])
            return body
        except Exception:
            return ""

    scored: list[tuple[int, dict, str]] = []
    with ThreadPoolExecutor(max_workers=CHEAP_POOL_WORKERS) as ex:
        futs = {ex.submit(_quick_sniff, it): it for it in scope_items}
        for fut in as_completed(futs):
            it = futs[fut]
            try:
                body = fut.result() or ""
            except Exception:
                body = ""
            if body:
                s = _score_in_text(body, name_terms)
                if s > 0:
                    scored.append((s, it, body))
            if time.perf_counter() - t0 > TIME_BUDGET:
                break

    # 체크리스트 스코프에서 득점 0이면 즉시 부정
    if not scored and strict_scope:
        return _kv(msg=f"선택한 폴더 범위 내에서 '{query}' 관련 파일을 찾지 못했습니다.")

    # 스코프 미설정일 때만 깊은 폴백 유지
    if not scored and (not strict_scope) and (time.perf_counter() - t0) <= FALLBACK_TIME_BUDGET_SEC:
        deep_cap = DEEP_PDF_MAX if no_name_signal else DEEP_PDF_MAX_HINT
        deep_targets = [it for it in scope_items if (it.get("ext") or "").lower() == ".pdf"][:deep_cap]
        for it in deep_targets:
            if time.perf_counter() - t0 > FALLBACK_TIME_BUDGET_SEC:
                break
            body = _fulltext_for_scoring(it, pdf_pages=12)
            if body:
                s = _score_in_text(body, name_terms)
                if s > 0:
                    scored.append((s, it, body))

    if not scored:
        return _kv(msg=f"'{query}' 관련 파일을 찾지 못했습니다.")

    scored.sort(key=lambda x: x[0], reverse=True)
    best_item = scored[0][1]
    try:
        preview = resolve_file_handler(best_item)
    except Exception:
        preview = "⚠️ 후보 파일 미리보는 중 오류가 발생했습니다."
    url = f"https://{_UI_HOST}/preview-pages/drives/{best_item['id']}"

    return _kv(best_item.get("name",""), best_item.get("_path","/"),
               best_item.get("ext",""), url, preview, "")














